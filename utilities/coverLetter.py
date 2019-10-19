import os
from os import path
import sys
import docx
import requests
from docx.shared import Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH
from bs4 import BeautifulSoup

inputDoc = docx.Document('inputs/input-cover-letter.docx')

def getCompanyAndTitle(url):
    page = requests.get(url)
    soup = BeautifulSoup(page.content, 'html.parser')

    company = soup.find("span", {"class": "company-name"}).getText()
    company = company.replace('at ', '').strip()
    # removed because I'm only applying for software engineering roles
    # title = soup.find("h1", {"class": "app-title"}).getText()
    return (company, 'software engineer')

def generateCoverLetter(url):
    # Options for word document formatting
    outputDoc = docx.Document()
    style = outputDoc.styles['Normal']
    font = style.font
    font.name = 'Times New Roman'
    font.size = Pt(12)

    company, title = getCompanyAndTitle(url)

    for i in range(len(inputDoc.paragraphs)-1):
        # Replace filler information
        text = inputDoc.paragraphs[i].text
        if '<company>' in text:
            text = text.replace('<company>', company)
        if '<title>' in text:
            text = text.replace('<title>', title)
        paragraph = outputDoc.add_paragraph(text)
        paragraph_format = paragraph.paragraph_format

        # Center first three paragraphs
        if i < 4:
            paragraph_format.alignment = WD_ALIGN_PARAGRAPH.CENTER

        # Add paragraph formatting
        paragraph_format.space_after = Pt(3)
        paragraph_format.space_after = Pt(3)
        paragraph_format.line_spacing = 1.0
        paragraph.style = outputDoc.styles['Normal']

    outputDoc.save('outputs/CoverLetter.docx')